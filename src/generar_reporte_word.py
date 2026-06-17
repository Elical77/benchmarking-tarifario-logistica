import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


MES_RELEVAMIENTO = "2026-05"
MES_TEXTO = "Mayo 2026"

ARCHIVO_COMPARATIVA = "comparativa_precios.xlsx"
ARCHIVO_FADEEAC = "fadeeac_indices.xlsx"

SALIDA_DOCX = f"reporte_comparativa_{MES_RELEVAMIENTO}.docx"
GRAFICO_FADEEAC = f"grafico_fadeeac_{MES_RELEVAMIENTO}.png"


def formatear_pct(valor):
    if pd.isna(valor):
        return "s/d"
    return f"{valor * 100:.1f}%".replace(".", ",")


def formatear_pesos(valor):
    if pd.isna(valor):
        return "s/d"
    return f"${valor:,.0f}".replace(",", ".")


def normalizar_tasa(valor):
    if pd.isna(valor):
        return None

    if isinstance(valor, str):
        txt = valor.strip().replace("%", "").replace(",", ".")
        return float(txt) / 100

    valor = float(valor)

    if valor > 1:
        return valor / 100

    return valor


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text_color(cell, color_hex):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor.from_string(color_hex)


def agregar_titulo(documento, texto, nivel=1):
    p = documento.add_paragraph()
    run = p.add_run(texto)
    run.bold = True

    if nivel == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(31, 78, 121)
    else:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0, 0, 0)

    return p


def agregar_parrafo(documento, texto):
    p = documento.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(texto)
    run.font.size = Pt(10)
    return p


def agregar_bullet(documento, texto):
    p = documento.add_paragraph(style="List Bullet")
    run = p.add_run(texto)
    run.font.size = Pt(10)
    return p


def leer_comparativa():
    df = pd.read_excel(ARCHIVO_COMPARATIVA, sheet_name="Comparativa Mes", header=2)
    df = df[df["destino"].notna()].copy()
    return df


def leer_fadeeac():
    df = pd.read_excel(ARCHIVO_FADEEAC)

    df.columns = [str(c).strip().lower() for c in df.columns]

    if "mes" not in df.columns or "indice" not in df.columns:
        raise Exception("fadeeac_indices.xlsx debe tener columnas: mes | indice")

    df["mes_fecha"] = pd.to_datetime(df["mes"], errors="coerce", dayfirst=True)
    df["indice_num"] = df["indice"].apply(normalizar_tasa)

    df = df.dropna(subset=["mes_fecha", "indice_num"]).copy()
    df = df.sort_values("mes_fecha")

    return df


def obtener_fadeeac_mes(df_fadeeac):
    periodo = pd.Period(MES_RELEVAMIENTO, freq="M")
    df_fadeeac["periodo"] = df_fadeeac["mes_fecha"].dt.to_period("M")

    fila = df_fadeeac[df_fadeeac["periodo"] == periodo]

    if fila.empty:
        return None

    return float(fila.iloc[0]["indice_num"])


def obtener_fadeeac_mes_anterior(df_fadeeac):
    periodo = pd.Period(MES_RELEVAMIENTO, freq="M") - 1
    df_fadeeac["periodo"] = df_fadeeac["mes_fecha"].dt.to_period("M")

    fila = df_fadeeac[df_fadeeac["periodo"] == periodo]

    if fila.empty:
        return None

    return float(fila.iloc[0]["indice_num"])


def generar_grafico_fadeeac(df_fadeeac):
    df_plot = df_fadeeac.tail(13).copy()
    df_plot["label"] = df_plot["mes_fecha"].dt.strftime("%b-%y")

    plt.figure(figsize=(8, 3.5))
    plt.plot(df_plot["label"], df_plot["indice_num"] * 100, marker="o")
    plt.title("Evolución mensual del Índice FADEEAC")
    plt.ylabel("Variación mensual (%)")
    plt.xticks(rotation=45, ha="right")
    plt.grid(axis="y", alpha=0.3)
    plt.tight_layout()
    plt.savefig(GRAFICO_FADEEAC, dpi=180)
    plt.close()


def calcular_resumen_general(df_ok):
    competidores = [
        ("Vía Cargo", "precio_via_cargo", "desvio_vc_%"),
        ("Andreani", "precio_andreani", "desvio_andreani_%"),
        ("OCA", "precio_oca", "desvio_oca_%"),
    ]

    registros = []

    for empresa, col_precio, col_desvio in competidores:
        tmp = df_ok[[col_precio, col_desvio]].dropna()

        if tmp.empty:
            continue

        desvio = tmp[col_desvio].mean()

        if desvio < 0:
            lectura = "Más barato que Cruz del Sur"
        elif desvio > 0:
            lectura = "Más caro que Cruz del Sur"
        else:
            lectura = "En línea con Cruz del Sur"

        registros.append({
            "empresa": empresa,
            "precio_promedio": tmp[col_precio].mean(),
            "desvio_promedio": desvio,
            "lectura": lectura,
            "casos": len(tmp),
        })

    return pd.DataFrame(registros)


def calcular_resumen_destinos(df_ok):
    resumen = df_ok.groupby("destino").agg(
        casos=("kg", "count"),
        cds_promedio=("precio_cruz_del_sur", "mean"),
        via_cargo_promedio=("precio_via_cargo", "mean"),
        andreani_promedio=("precio_andreani", "mean"),
        oca_promedio=("precio_oca", "mean"),
        desvio_vc=("desvio_vc_%", "mean"),
        desvio_andreani=("desvio_andreani_%", "mean"),
        desvio_oca=("desvio_oca_%", "mean"),
    ).reset_index()

    return resumen


def agregar_tabla_resumen_general(documento, resumen):
    tabla = documento.add_table(rows=1, cols=5)
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabla.style = "Table Grid"

    headers = ["Competidor", "Precio prom.", "Dif. vs CDS", "Lectura", "Casos"]

    for i, h in enumerate(headers):
        cell = tabla.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "1F4E78")
        set_cell_text_color(cell, "FFFFFF")

    for _, row in resumen.iterrows():
        cells = tabla.add_row().cells
        cells[0].text = row["empresa"]
        cells[1].text = formatear_pesos(row["precio_promedio"])
        cells[2].text = formatear_pct(row["desvio_promedio"])
        cells[3].text = row["lectura"]
        cells[4].text = str(int(row["casos"]))

    return tabla


def agregar_tabla_destinos(documento, resumen_destinos):
    tabla = documento.add_table(rows=1, cols=5)
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabla.style = "Table Grid"

    headers = ["Destino", "CDS prom.", "Vía Cargo", "Andreani", "OCA"]

    for i, h in enumerate(headers):
        cell = tabla.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "1F4E78")
        set_cell_text_color(cell, "FFFFFF")

    for _, row in resumen_destinos.iterrows():
        cells = tabla.add_row().cells
        cells[0].text = row["destino"]
        cells[1].text = formatear_pesos(row["cds_promedio"])
        cells[2].text = formatear_pct(row["desvio_vc"])
        cells[3].text = formatear_pct(row["desvio_andreani"])
        cells[4].text = formatear_pct(row["desvio_oca"])

    return tabla


def texto_competidor(resumen_general, empresa, descripcion_servicio):
    fila = resumen_general[resumen_general["empresa"] == empresa]

    if fila.empty:
        return f"No se obtuvieron valores válidos para {empresa} en el relevamiento de {MES_TEXTO}."

    r = fila.iloc[0]
    desvio = r["desvio_promedio"]

    if desvio < 0:
        comparacion = "inferiores"
    elif desvio > 0:
        comparacion = "superiores"
    else:
        comparacion = "en línea"

    return (
        f"Durante {MES_TEXTO}, los precios de {empresa} resultaron en promedio "
        f"{formatear_pct(abs(desvio))} {comparacion} respecto de Cruz del Sur, "
        f"considerando los destinos y rangos de kilogramos relevados. "
        f"Para la comparación se tomó como referencia {descripcion_servicio}."
    )


def crear_reporte():
    df = leer_comparativa()
    df_ok = df[df["estado_cruce"] == "OK"].copy()
    df_no_ok = df[df["estado_cruce"] != "OK"].copy()

    df_fadeeac = leer_fadeeac()
    tasa_fadeeac = obtener_fadeeac_mes(df_fadeeac)
    tasa_fadeeac_anterior = obtener_fadeeac_mes_anterior(df_fadeeac)
    generar_grafico_fadeeac(df_fadeeac)

    resumen_general = calcular_resumen_general(df_ok)
    resumen_destinos = calcular_resumen_destinos(df_ok)

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    # Portada simple y editable
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CRUZ DEL SUR")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(31, 78, 121)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Comparativa de mercado")
    run.bold = True
    run.font.size = Pt(28)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(MES_TEXTO)
    run.font.size = Pt(16)

    doc.add_page_break()

    # Índice textual
    agregar_titulo(doc, "Contenido", nivel=1)
    agregar_parrafo(doc, "Capítulo 1 – Situación del Mercado")
    agregar_parrafo(doc, "Índice de Costos de Transporte – FADEEAC")
    agregar_parrafo(doc, "Capítulo 2 – Relevamiento")
    agregar_parrafo(doc, "Capítulo 2.1 – Andreani")
    agregar_parrafo(doc, "Capítulo 2.2 – OCA")
    agregar_parrafo(doc, "Capítulo 2.3 – Vía Cargo")
    agregar_parrafo(doc, "Aclaraciones metodológicas")
    agregar_parrafo(doc, "Referencias")

    doc.add_page_break()

    # Capítulo 1
    agregar_titulo(doc, "Capítulo 1 – Situación del Mercado", nivel=1)
    agregar_titulo(doc, "Índice de Costos de Transporte – FADEEAC", nivel=2)

    if tasa_fadeeac is not None:
        texto_fadeeac = (
            f"El Índice de Costos de Transporte de la FADEEAC registró en {MES_TEXTO} "
            f"una variación mensual de {formatear_pct(tasa_fadeeac)}. "
        )

        if tasa_fadeeac_anterior is not None:
            if tasa_fadeeac < tasa_fadeeac_anterior:
                texto_fadeeac += (
                    f"El dato muestra una desaceleración respecto del mes anterior "
                    f"({formatear_pct(tasa_fadeeac_anterior)}). "
                )
            else:
                texto_fadeeac += (
                    f"El dato se ubicó por encima del mes anterior "
                    f"({formatear_pct(tasa_fadeeac_anterior)}). "
                )

        texto_fadeeac += (
            "De acuerdo con la información publicada por FADEEAC, los costos del transporte "
            "acumulan aproximadamente 20% en lo que va de 2026. Si bien la variación mensual "
            "de mayo fue menor que la de marzo y abril, el sector continúa operando con presión "
            "de costos, especialmente por el comportamiento del combustible, los gastos generales, "
            "la mano de obra y el deterioro de la infraestructura vial."
        )

        agregar_parrafo(doc, texto_fadeeac)

        agregar_parrafo(
            doc,
            "Este contexto resulta relevante para interpretar los movimientos tarifarios del mercado: "
            "una desaceleración mensual del índice no implica una normalización plena de costos, sino "
            "una moderación respecto de meses previos de mayor presión."
        )

    else:
        agregar_parrafo(
            doc,
            f"No se encontró un valor de FADEEAC para {MES_TEXTO} en el archivo de índices."
        )

    doc.add_picture(GRAFICO_FADEEAC, width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Capítulo 2
    agregar_titulo(doc, "Capítulo 2 – Relevamiento", nivel=1)

    agregar_parrafo(
        doc,
        "La siguiente información fue relevada a partir de cotizadores online, tarifarios publicados "
        "y tarifas internas de Cruz del Sur. La comparación considera servicios de entrega en sucursal "
        "o agencia, según disponibilidad de cada operador."
    )

    if not df_no_ok.empty:
        destinos_no_comparables = ", ".join(sorted(df_no_ok["destino"].dropna().unique()))
        agregar_parrafo(
            doc,
            f"Se excluyeron de la comparación principal aquellos casos sin tarifa disponible de Cruz del Sur "
            f"o sin valor válido de algún competidor. En esta corrida, los destinos/casos observados fuera "
            f"del cruce completo fueron: {destinos_no_comparables}."
        )

    agregar_titulo(doc, "Resumen general", nivel=2)
    agregar_tabla_resumen_general(doc, resumen_general)

    agregar_titulo(doc, "Resumen por destino", nivel=2)
    agregar_tabla_destinos(doc, resumen_destinos)

    agregar_titulo(doc, "Capítulo 2.1 – Andreani", nivel=2)
    agregar_parrafo(
        doc,
        texto_competidor(
            resumen_general,
            "Andreani",
            "el precio de envío a sucursal informado por el cotizador online, neto de IVA"
        )
    )
    agregar_parrafo(
        doc,
        "El cotizador de Andreani informa valores con IVA incluido. Para mantener la comparabilidad "
        "con Cruz del Sur, los valores utilizados en el análisis fueron neteados de IVA dividiendo "
        "el precio publicado por 1,21. Asimismo, el cotizador indica que para proteger el envío "
        "se adiciona un 1% del valor declarado."
    )

    agregar_titulo(doc, "Capítulo 2.2 – OCA", nivel=2)
    agregar_parrafo(
        doc,
        texto_competidor(
            resumen_general,
            "OCA",
            "la tarifa de entrega en sucursal del tarifario vigente, neta de IVA"
        )
    )
    agregar_parrafo(
        doc,
        "La clasificación zonal de OCA se realizó tomando como origen Buenos Aires y considerando las "
        "categorías Local, Regional, Nacional 1 y Nacional 2. Para el cálculo se utilizó el servicio "
        "de entrega en sucursal, por ser el más comparable con la modalidad de Cruz del Sur. "
        "Dado que el tarifario de OCA informa precios con IVA incluido, los valores utilizados en "
        "la comparación fueron neteados de IVA dividiendo el precio publicado por 1,21."
    )

    agregar_titulo(doc, "Capítulo 2.3 – Vía Cargo", nivel=2)
    agregar_parrafo(
        doc,
        texto_competidor(
            resumen_general,
            "Vía Cargo",
            "el servicio despacho agencia – entrega agencia"
        )
    )
    agregar_parrafo(
        doc,
        "El relevamiento fue realizado mediante cotizador online. Las dimensiones utilizadas fueron calculadas "
        "para evitar que el peso volumétrico supere el peso real del envío. Para mantener la comparabilidad, "
        "se tomó el valor correspondiente a despacho en agencia y entrega en agencia."
    )

    agregar_titulo(doc, "Aclaraciones metodológicas", nivel=1)

    aclaraciones = [
        "El análisis se basa en precios finales publicados o tarifarios vigentes y no contempla descuentos comerciales, promociones ni acuerdos particulares.",
        "Para Cruz del Sur se utilizó el tarifario contado vigente del mes relevado.",
        "Para Vía Cargo se tomó el servicio despacho agencia – entrega agencia.",
        "Para Andreani se tomó el precio de envío a sucursal informado por el cotizador online, neto de IVA.",
        "Para OCA se tomó la tarifa de entrega en sucursal del tarifario publicado, neta de IVA.",
        "Andreani y OCA publican precios con IVA incluido; para la comparación se dividieron dichos valores por 1,21.",
        "El relevamiento corresponde a mayo 2026.",
        "Los valores pueden presentar diferencias respecto de una operación real por criterios de admisión, seguros, valor declarado, cargos adicionales o validación en mostrador.",
    ]

    for item in aclaraciones:
        agregar_bullet(doc, item)

    agregar_titulo(doc, "Referencias", nivel=1)

    referencias = [
        "FADEEAC – Índice de Costos del Transporte.",
        "Andreani – Cotizador online.",
        "OCA – Lista de precios OCA Express Pak.",
        "Vía Cargo – Cotizador online.",
        "Cruz del Sur – Tarifario contado.",
    ]

    for ref in referencias:
        agregar_bullet(doc, ref)

    doc.save(SALIDA_DOCX)

    print("Reporte generado correctamente.")
    print(f"Archivo: {SALIDA_DOCX}")


if __name__ == "__main__":
    crear_reporte()